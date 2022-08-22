import docx


def save_to_cat_table(cat):
    doc = docx.Document()
    table = doc.add_table(rows=1, cols=3)

    doc.add_heading('Cat Table', 0)

    row = table.rows[0].cells
    row[0].text = 'Id'
    row[1].text = 'Name'
    row[2].text = 'Status'

    for cat_id, cat_name, status in cat:
        row = table.add_row().cells
        row[0].text = cat_id
        row[1].text = cat_name
        row[2].text = status

    doc.save('cat.docx')
